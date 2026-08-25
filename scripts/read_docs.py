import sys, os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from docx import Document

files = {
    'REPORT': r'c:\Users\sange\OneDrive\Desktop\Physics-Informed-Quantum-Reservoir-Transformer\docs\FINAL_RESEARCH_REPORT_COMPLETE.docx',
    'MANUSCRIPT': r'c:\Users\sange\OneDrive\Desktop\Physics-Informed-Quantum-Reservoir-Transformer\docs\Elsevier_Q1_Manuscript_Draft.docx',
}

for label, fpath in files.items():
    print('='*70)
    print(f'DOCUMENT: {label}')
    print('='*70)
    doc = Document(fpath)
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            print(t)
    for ti, table in enumerate(doc.tables):
        print(f'\n[TABLE {ti+1}]')
        for row in table.rows:
            print(' | '.join(c.text.strip() for c in row.cells))
    print()
